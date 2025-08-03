import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
import tempfile
import os

def rimuovi_sfondo_riga(riga):
    for cell in riga.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            tcPr.remove(shd)

def genera_scheda(df, template_bytes):
    df_sorted = df.sort_values(by="Label")
    document = Document(template_bytes)
    tabella_4 = document.tables[3]
    tabella_5 = document.tables[4]

    if len(tabella_4.rows) > 2:
        rimuovi_sfondo_riga(tabella_4.rows[2])
    while len(tabella_4.rows) > 2:
        tabella_4._tbl.remove(tabella_4.rows[2]._tr)

    nc_counter = 1
    oss_counter = 1
    doc_to_has_nc = defaultdict(bool)
    general_issues = []
    document_issues = defaultdict(list)

    for _, row in df_sorted.iterrows():
        tag = str(row["Tags"]).strip().upper()
        if tag not in ["NC", "OSS"]:
            continue
        codice_doc = row["Title"]
        descrizione = row["Description"]
        ispettore = row["Created by"]
        if "rilievi generali" in str(codice_doc).lower():
            general_issues.append((tag, codice_doc, descrizione, ispettore))
        else:
            document_issues[codice_doc].append((tag, codice_doc, descrizione, ispettore))

    def inserisci_riga(tag, codice_doc, descrizione, ispettore):
        nonlocal nc_counter, oss_counter
        codice_nc = f"{tag}{nc_counter if tag == 'NC' else oss_counter}"
        if tag == "NC":
            nc_counter += 1
        else:
            oss_counter += 1
        riga = tabella_4.add_row().cells
        run0 = riga[0].paragraphs[0].add_run(codice_nc)
        run0.bold = True
        run1 = riga[1].paragraphs[0].add_run(codice_doc)
        run1.bold = True
        riga[2].text = descrizione
        riga[3].text = ispettore

    for entry in general_issues:
        tag, codice_doc, descrizione, ispettore = entry
        doc_to_has_nc[codice_doc] = True
        inserisci_riga(tag, codice_doc, descrizione, ispettore)

    for codice_doc, entries in document_issues.items():
        for tag, _, descrizione, ispettore in entries:
            doc_to_has_nc[codice_doc] = True
            inserisci_riga(tag, codice_doc, descrizione, ispettore)

    while len(tabella_5.rows) > 1:
        tabella_5._tbl.remove(tabella_5.rows[1]._tr)

    for doc_title in df["Title"].dropna().unique():
        riga = tabella_5.add_row().cells
        run_doc = riga[0].paragraphs[0].add_run(doc_title)
        run_doc.bold = True
        if doc_to_has_nc[doc_title]:
            run_x = riga[3].paragraphs[0].add_run("X")
            run_x.bold = True
        else:
            run_x = riga[4].paragraphs[0].add_run("X")
            run_x.bold = True

    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    document.save(output_path.name)
    return output_path.name

st.set_page_config(page_title="Generatore Schede Ispettive", page_icon="📝")
st.title("📝 Generatore Schede Ispettive")

uploaded_excel = st.file_uploader("Carica il file ToDo (Excel)", type=["xlsx"])
uploaded_template = st.file_uploader("Carica il template Word", type=["docx"])

if uploaded_excel and uploaded_template:
    df = pd.read_excel(uploaded_excel)
    if st.button("Genera Scheda Ispettiva"):
        with st.spinner("Generazione in corso..."):
            output_docx = genera_scheda(df, uploaded_template)
            with open(output_docx, "rb") as f:
                st.download_button("📥 Scarica Scheda Ispettiva", f, file_name="Scheda_Ispettiva.docx")
            os.remove(output_docx)
